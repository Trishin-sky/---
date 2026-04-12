import streamlit as st
import pandas as pd
from datetime import datetime
import os
import json
from docx import Document
from docx.shared import Inches, Pt
from io import BytesIO
import base64
from pathlib import Path

# ==================== НАСТРОЙКИ СТРАНИЦЫ ====================
st.set_page_config(
    page_title="Астма-тест",
    page_icon= "https://img.icons8.com/color/96/000000/lungs.png",
    layout="wide"
)

# ==================== ИНИЦИАЛИЗАЦИЯ СЕССИИ ====================
if 'patients_db' not in st.session_state:
    # Загружаем существующую базу если есть
    if os.path.exists("patients_data.json"):
        with open("patients_data.json", "r", encoding="utf-8") as f:
            st.session_state.patients_db = json.load(f)
    else:
        st.session_state.patients_db = []

if 'current_patient' not in st.session_state:
    st.session_state.current_patient = {
        'id': None,
        'fio': '',
        'birth_date': None,
        'gender': '',
        'test_date': None,
        'act_answers': {},
        'act_score': None,
        'hads_a_answers': {},
        'hads_a_score': None,
        'hads_d_answers': {},
        'hads_d_score': None,
        'cirs_answers': {},
        'cirs_score': None,
        'completed_tests': []
    }

if 'page' not in st.session_state:
    st.session_state.page = "patient_info"

# ==================== ФУНКЦИИ ДЛЯ РАСЧЁТОВ ====================
def calculate_act(answers):
    """Расчёт ACT (5 вопросов, каждый от 1 до 5)"""
    if len(answers) < 5:
        return None
    total = sum(answers.values())
    return total

def interpret_act(score):
    if score is None:
        return "Не заполнено", "info"
    if score >= 20:
        return f"Хорошо контролируемая астма ({score} баллов)", "success"
    elif score >= 16:
        return f"Частично контролируемая астма ({score} баллов)", "warning"
    else:
        return f"Неконтролируемая астма ({score} баллов)", "error"

def calculate_hads(answers_anxiety, answers_depression):
    """Расчёт HADS (7 вопросов тревога + 7 вопросов депрессия)"""
    anxiety_score = sum(answers_anxiety.values()) if len(answers_anxiety) == 7 else None
    depression_score = sum(answers_depression.values()) if len(answers_depression) == 7 else None
    return anxiety_score, depression_score

def interpret_hads(score, scale_type):
    if score is None:
        return "Не заполнено", "info"
    if score <= 7:
        return f"Норма ({score} баллов)", "success"
    elif score <= 10:
        return f"Субклинические проявления ({score} баллов)", "warning"
    else:
        return f"Клинически выраженные симптомы ({score} баллов)", "error"

def calculate_cirs(answers):
    """Расчёт CIRS (14 систем органов)"""
    if len(answers) < 14:
        return None
    total = sum(answers.values())
    return total

def interpret_cirs(score):
    if score is None:
        return "Не заполнено", "info"
    if score == 0:
        return f"Коморбидная нагрузка отсутствует ({score} баллов)", "success"
    else:
        return f"Коморбидная нагрузка ({score} из 56 баллов)", "error"

# ==================== ВОПРОСЫ ДЛЯ ТЕСТОВ ====================
ACT_QUESTIONS = {
    1: "Как часто за последние 4 недели астма мешала Вам выполнять обычный объём работы в учебном заведении, на работе или дома?",
    2: "Как часто за последние 4 недели Вы отмечали у себя затрудненное дыхание?",
    3: "Как часто за последние 4 недели Вы просыпались ночью или раньше, чем обычно, из-за симптомов астмы (свистящего дыхания, кашля, затрудненного дыхания, чувства стеснения в груди или боли в груди)?",
    4: "Как часто за последние 4 недели Вы использовали быстродействующий ингалятор (например, Вентолин, Беротек, Беродуал, Атровепт, Сальбутамол, Самол, Сальбен, Астмопент) или небулайзер (аэрозольный аппарат) с лекарством (например, Беротек, Беродуал, Вентолин Небулы)?",
    5: "Как бы Вы оценили, насколько Вам удавалось контролировать астму за последние 4 недели??"
}

ACT_OPTIONS = {
    1: {1: "Все время", 2: "Очень часто", 3: "Всегда", 4: "Редко", 5: "Никогда"},
    2: {1: "Чаще, чем раз в день", 2: "Раз в день", 3: "От 3 до 6 раз в неделю", 4: "1-2 раза в неделю", 5: "Ни разу"},
    3: {1: "4 ночи в неделю или чаще", 2: "2-3 ночи в неделю", 3: "Раз в неделю", 4: "1 или 2 раза", 5: "Ни разу"},
    4: {1: "3 раза в день или чаще", 2: "1 или 2 раза в день", 3: "2 или 3 раза в неделю", 4: "Один раз в неделю или реже", 5: "Ни разу"},
    5: {1: "Совсем не удавалось контролировать", 2: "Плохо удавалось контролировать", 3: "В некоторой степени удавалось контролировать", 4: "Хорошо удавалось контролировать", 5: "Полностью удавалось контролировать"}
}

HADS_ANXIETY_QUESTIONS = {
    1: "Я испытываю напряжение, мне не по себе",
    2: "Я испытываю страх, кажется, что что-то ужасное может вот-вот случиться",
    3: "Беспокойные мысли крутятся у меня в голове",
    4: "Я легко могу присесть и расслабиться",
    5: "Я испытываю внутреннее напряжение или дрожь",
    6: "Я испытываю неусидчивость, мне постоянно нужно двигаться",
    7: "У меня бывает внезапное чувство паники"
}

HADS_DEPRESSION_QUESTIONS = {
    1: "То, что приносило мне большое удовольствие, и сейчас вызывает у меня такое же чувство",
    2: "Я способен рассмеяться и увидеть в том или ином событии смешное",
    3: "Я испытываю бодрость",
    4: "Мне кажется, что я стал всё делать очень медленно",
    5: "Я не слежу за своей внешностью",
    6: "Я считаю, что мои дела (занятия, увлечения) могут принести мне чувство удовлетворения",
    7: "Я могу получить удовольствие от хорошей книги, радио- или телепрограммы"
}

HADS_ANXIETY_OPTIONS = {
    1: {3: "Все время", 2: "Часто", 1: "Время от времени, иногда", 0: "Совсем не испытываю"},
    2: {3: "Определённо это так, и страх очень велик", 2: "Да, это так, но страх не очень велик", 1: "Иногда, но это меня не беспокоит", 0: "Совсем не испытываю"},
    3: {3: "Постоянно", 2: "Большую часть времени", 1: "Время от времени и не так часто", 0: "Только иногда"},
    4: {0: "Определённо, это так", 1: "Наверное, это так", 2: "Лишь изредка, это так", 3: "Совсем не могу"},
    5: {0: "Совсем не испытываю", 1: "Иногда", 2: "Часто", 3: "Очень часто"},
    6: {3: "Определённо, это так", 2: "Наверное, это так", 1: "Лишь в некоторой степени, это так", 0: "Совсем не испытываю"},
    7: {3: "Очень часто", 2: "Довольно часто", 1: "Не так уж часто", 0: "Совсем не бывает"},
}

HADS_DEPRESSION_OPTIONS = {
    1: {0: "Определённо, это так", 1: "Наверное, это так", 2: "Лишь в очень малой степени, это так", 3: "Это совсем не так"},
    2: {0: "Определённо, это так", 1: "Наверное, это так", 2: "Лишь в очень малой степени, это так", 3: "Совсем не способен"},
    3: {3: "Совсем не испытываю", 2: "Очень редко", 1: "Иногда", 0: "Практически всё время"},
    4: {3: "Практически все время", 2: "Часто", 1: "Иногда", 0: "Совсем нет"},
    5: {3: "Определённо, это так", 2: "Я не уделяю этому столько времени, сколько нужно", 1: "Может быть, я стал меньше уделять этому времени", 0: "Я слежу за собой так же, как и раньше"},
    6: {0: "Точно так же, как и обычно", 1: "Да, но не в той степени, как раньше", 2: "Значительно меньше, чем обычно", 3: "Совсем так не считаю"},
    7: {0: "Часто", 1: "Иногда", 2: "Редко", 3: "Очень редко"},
}


CIRS_SYSTEMS = {
    1: "Болезни сердца",
    2: "Болезни сосудов (кровеносных и лимфатических)",
    3: "Болезни крови (костного мозга, селезенки и периферической крови)",
    4: "Болезни органов дыхательной системы (трахеи, бронхов и легких)",
    5: "Болезни органов чувств (глаз, носа, ушей, глотки и гортани)",
    6: "Болезни органов верхних отделов пищеварительной системы (пищевода, желудка, двенадцатиперстной кишки, поджелудочной железы (не включая СД) и желчного пузыря)",
    7: "Болезни органов нижних отделов пищеварительной системы (тонкого и толстого кишечника)",
    8: "Болезни печени",
    9: "Болезни почек",
    10: "Болезни органов мочеполовой системы (мочеточников, мочевого пузыря, мочеиспускательного канала, предстательной железы и половых органов)",
    11: "Болезни органов опорно-двигательной системы (мышц, суставов, костей) и кожных покровов",
    12: "Болезни органов центральной и периферической нервной системы (головного мозга, спинного мозга и нервов)",
    13: "Болезни органов эндокринной системы и нарушения метаболизма (включая сахарный диабет)",
    14: "Психические нарушения"
}

CIRS_OPTIONS = {
    0: "0 — Отсутствие заболеваний в этой системе органов или наличие патологии, которая не мешает нормальной жизнедеятельности, не влияет на прогноз и не требует лечения",
    1: "1 — Легкие отклонения от нормы или перенесенные в прошлом заболевания",
    2: "2 — Заболевание, при котором необходимо назначение медикаментозной терапии",
    3: "3 — Заболевание, ставшее причиной инвалидности",
    4: "4 — Жизнеугрожающее заболевание, требующее проведения неотложной терапии"
}

# ==================== ФУНКЦИЯ СОЗДАНИЯ WORD-ОТЧЁТА ====================
def create_word_report(patient, act_score, act_interpretation, hads_a_score, hads_a_int, 
                       hads_d_score, hads_d_int, cirs_score, cirs_int, test_date):
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('Астма-тест — Результаты скрининга', 0)
    title.alignment = 1  # Центр
    
    # Информация о пациенте
    doc.add_heading('Информация о пациенте', level=1)
    doc.add_paragraph(f'ФИО: {patient["fio"]}')
    doc.add_paragraph(f'Дата рождения: {patient["birth_date"]}')
    doc.add_paragraph(f'Пол: {patient["gender"]}')
    doc.add_paragraph(f'Дата тестирования: {test_date}')
    
    
    
    # Результаты ACT
    doc.add_heading('1. ACT (Asthma Control Test) — контроль бронхиальной астмы', level=1)
    doc.add_paragraph(f'Результат: {act_interpretation}')
    doc.add_paragraph('Интерпретация:')
    doc.add_paragraph('• 20-25 баллов — хорошо контролируемая астма', style='List Bullet')
    doc.add_paragraph('• 16-19 баллов — частично контролируемая астма', style='List Bullet')
    doc.add_paragraph('• 15 и менее — неконтролируемая астма', style='List Bullet')
    
   
    
    # Результаты HADS
    doc.add_heading('2. HADS (Hospital Anxiety and Depression Scale)', level=1)
    doc.add_paragraph(f'Тревога (HADS-A): {hads_a_int}')
    doc.add_paragraph(f'Депрессия (HADS-D): {hads_d_int}')
    doc.add_paragraph('Интерпретация:')
    doc.add_paragraph('• 0-7 баллов — норма', style='List Bullet')
    doc.add_paragraph('• 8-10 баллов — субклинические проявления', style='List Bullet')
    doc.add_paragraph('• 11+ баллов — клинически выраженные симптомы', style='List Bullet')
    
    
    
    # Результаты CIRS
    doc.add_heading('3. CIRS (Cumulative Illness Rating Scale) — коморбидность', level=1)
    doc.add_paragraph(f'Результат: {cirs_int}')
    
    # Подвал
    doc.add_paragraph()
    doc.add_paragraph(f'Отчёт сгенерирован автоматически системой Астма-тест v1.0')
    
    return doc


def create_excel_report(patient):
    """Создаёт Excel-файл с данными текущего пациента и возвращает его для скачивания"""
    # Создаём DataFrame с данными текущего пациента
    data = [{
        'ID': patient['id'],
        'ФИО': patient['fio'],
        'Дата рождения': patient['birth_date'],
        'Пол': patient['gender'],
        'Дата тестирования': patient['test_date'],
        'ACT (баллы)': patient['act_score'],
        'ACT (интерпретация)': interpret_act(patient['act_score'])[0],
        'HADS-Тревога (баллы)': patient['hads_a_score'],
        'HADS-Тревога (интерпретация)': interpret_hads(patient['hads_a_score'], "anxiety")[0],
        'HADS-Депрессия (баллы)': patient['hads_d_score'],
        'HADS-Депрессия (интерпретация)': interpret_hads(patient['hads_d_score'], "depression")[0],
        'CIRS (баллы)': patient['cirs_score'],
        'CIRS (интерпретация)': interpret_cirs(patient['cirs_score'])[0]
    }]
    
    df = pd.DataFrame(data)
    
    # Сохраняем в BytesIO (виртуальный файл в памяти)
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Результаты пациента')
    
    output.seek(0)  # Возвращаем курсор в начало файла
    return output



# ==================== ФУНКЦИЯ СОХРАНЕНИЯ В БАЗУ ====================
def save_to_database(patient):
    # Проверяем, существует ли уже такой пациент (по ID)
    existing_ids = [p['id'] for p in st.session_state.patients_db]
    if patient['id'] in existing_ids:
        # Обновляем существующего
        for i, p in enumerate(st.session_state.patients_db):
            if p['id'] == patient['id']:
                st.session_state.patients_db[i] = patient.copy()
                break
    else:
        # Добавляем нового
        st.session_state.patients_db.append(patient.copy())
    
    # Сохраняем в JSON
    with open("patients_data.json", "w", encoding="utf-8") as f:
        json.dump(st.session_state.patients_db, f, ensure_ascii=False, indent=2)
    
    # Сохраняем в Excel
    save_to_excel(st.session_state.patients_db)

    excel_file = create_excel_report(patient)

def save_to_excel(patients_db):
    if not patients_db:
        return
    
    data = []
    for p in patients_db:
        row = {
            'ID': p['id'],
            'ФИО': p['fio'],
            'Дата рождения': p['birth_date'],
            'Пол': p['gender'],
            'Дата тестирования': p['test_date'],
            'ACT (баллы)': p['act_score'],
            'HADS-Тревога (баллы)': p['hads_a_score'],
            'HADS-Депрессия (баллы)': p['hads_d_score'],
            'CIRS (баллы)': p['cirs_score']
        }
        data.append(row)
    
    df = pd.DataFrame(data)
    df.to_excel("data/patients_database.xlsx", index=False)




# ==================== СТРАНИЦА: ИНФОРМАЦИЯ О ПАЦИЕНТЕ ====================
def render_patient_info():
    st.title("Астма-тест — Скрининг психоэмоциональных нарушений")
    st.markdown("### Шаг 1: Введите данные пациента")
    
    col1, col2 = st.columns(2)
    
    with col1:
        fio = st.text_input("ФИО пациента *", value=st.session_state.current_patient['fio'])
        birth_date = st.date_input("Дата рождения", value=None, min_value=datetime(1900, 1, 1))
    
    with col2:
        gender = st.radio("Пол", ["Мужской", "Женский"], horizontal=True)
    
    if st.button("✅ Начать тестирование", type="primary", use_container_width=True):
        if fio and birth_date:
            st.session_state.current_patient['id'] = datetime.now().strftime('%Y%m%d%H%M%S')
            st.session_state.current_patient['fio'] = fio
            st.session_state.current_patient['birth_date'] = birth_date.strftime('%d.%m.%Y')
            st.session_state.current_patient['gender'] = gender
            st.session_state.current_patient['test_date'] = datetime.now().strftime('%d.%m.%Y %H:%M')
            st.session_state.current_patient['completed_tests'] = []
            st.session_state.page = "act_test"
            st.rerun()
        else:
            st.error("Пожалуйста, заполните ФИО и дату рождения")

# ==================== ТЕСТ ACT ====================
def render_act_test():
    st.title("📋 Тест ACT (Asthma Control Test)")
    st.markdown("Оцените контроль бронхиальной астмы за **последние 4 недели**")
    
    answers = {}
    for q_num, question in ACT_QUESTIONS.items():
        st.markdown(f"**Вопрос {q_num}:** {question}")
        answer = st.radio(
            "Выберите ответ:",
            options=list(ACT_OPTIONS[q_num].keys()),
            format_func=lambda x: ACT_OPTIONS[q_num][x],
            key=f"act_q{q_num}",
            label_visibility="collapsed"
        )
        answers[q_num] = answer
        st.divider()
    
    col1, col2 = st.columns(2)
    
    if st.button("💾 Сохранить и перейти к HADS", type="primary", use_container_width=True):
        if len(answers) == 5:
            score = calculate_act(answers)
            st.session_state.current_patient['act_answers'] = answers
            st.session_state.current_patient['act_score'] = score
            if 'ACT' not in st.session_state.current_patient['completed_tests']:
                st.session_state.current_patient['completed_tests'].append('ACT')
                st.success(f"Результаты сохранены! Баллы: {score}")
                st.session_state.page = "hads_test"
                st.rerun()
        else:
            st.warning("Ответьте на все вопросы (5 вопросов)")

# ==================== ТЕСТ HADS ====================
def render_hads_test():
    st.title("📋 Тест HADS (Hospital Anxiety and Depression Scale)")
    st.markdown("Оцените, как Вы себя чувствовали **за последнюю неделю**")
    
    st.subheader("Часть 1 — Тревога")
    anxiety_answers = {}
    for q_num, question in HADS_ANXIETY_QUESTIONS.items():
        st.markdown(f"**{q_num}.** {question}")
        answer = st.radio(
            "Выберите ответ:",
            options=list(HADS_ANXIETY_OPTIONS[q_num].keys()),
            format_func=lambda x: HADS_ANXIETY_OPTIONS[q_num][x],
            key=f"hads_a_q{q_num}",
            label_visibility="collapsed"
        )
        anxiety_answers[q_num] = answer
        st.divider()
    
    st.subheader("Часть 2 — Депрессия")
    depression_answers = {}
    for q_num, question in HADS_DEPRESSION_QUESTIONS.items():
        st.markdown(f"**{q_num}.** {question}")
        answer = st.radio(
            "Выберите ответ:",
            options=list(HADS_DEPRESSION_OPTIONS[q_num].keys()),
            format_func=lambda x: HADS_DEPRESSION_OPTIONS[q_num][x],
            key=f"hads_d_q{q_num}",
            label_visibility="collapsed"
        )
        depression_answers[q_num] = answer
        st.divider()
    
    col1, col2 = st.columns(2)
    if st.button("💾 Сохранить результаты HADS и перейти к CIRS", type="primary", use_container_width=True):
        if len(anxiety_answers) == 7 and len(depression_answers) == 7:
            a_score, d_score = calculate_hads(anxiety_answers, depression_answers)
            st.session_state.current_patient['hads_a_answers'] = anxiety_answers
            st.session_state.current_patient['hads_a_score'] = a_score
            st.session_state.current_patient['hads_d_answers'] = depression_answers
            st.session_state.current_patient['hads_d_score'] = d_score
            if 'HADS' not in st.session_state.current_patient['completed_tests']:
                st.session_state.current_patient['completed_tests'].append('HADS')
                st.success(f"Результаты сохранены! Тревога: {a_score}, Депрессия: {d_score}")
                st.session_state.page = "cirs_test"
                st.rerun()
    else:
        st.warning("Ответьте на все вопросы (7 вопросов о тревоге и 7 о депрессии)")

# ==================== ТЕСТ CIRS ====================
def render_cirs_test():
    st.title("📋 Тест CIRS (Cumulative Illness Rating Scale)")
    st.markdown("Оцените состояние органов и систем")
    
    answers = {}
    for sys_num, system_name in CIRS_SYSTEMS.items():
        st.markdown(f"**Система {sys_num}:** {system_name}")
        answer = st.radio(
            "Выберите степень тяжести:",
            options=list(CIRS_OPTIONS.keys()),
            format_func=lambda x: CIRS_OPTIONS[x],
            key=f"cirs_q{sys_num}",
            label_visibility="collapsed"
        )
        answers[sys_num] = answer
        st.divider()
    
    col1, col2 = st.columns(2)
    if st.button("💾 Сохранить результаты CIRS и перейти к результатам", type="primary", use_container_width=True):
        if len(answers) == 14:
            score = calculate_cirs(answers)
            st.session_state.current_patient['cirs_answers'] = answers
            st.session_state.current_patient['cirs_score'] = score
            if 'CIRS' not in st.session_state.current_patient['completed_tests']:
                st.session_state.current_patient['completed_tests'].append('CIRS')
                st.success(f"Результаты сохранены! Баллы: {score}")
                st.session_state.page = "results"
                st.rerun()
    else:
        st.warning("Оцените все системы органов (14 систем)")

# ==================== СТРАНИЦА РЕЗУЛЬТАТОВ ====================
def render_results():
    st.title("📊 Результаты скрининга")
    
    patient = st.session_state.current_patient
    
    # Отображаем информацию о пациенте
    with st.expander("👤 Информация о пациенте", expanded=True):
        col1, col2, col3 = st.columns(3)
        col1.metric("ФИО", patient['fio'])
        col2.metric("Дата рождения", patient['birth_date'])
        col3.metric("Пол", patient['gender'])
        st.caption(f"Дата тестирования: {patient['test_date']}")

    st.markdown("""
<style>
    .st-emotion-cache-1qixc23 {
        /* Размеры */
        font-size: 18px;
        
        /* Отступы */
        padding: 0px;
        margin: 0px;
                
        /* Шрифты */
        font-family: 'Segoe UI', Arial, sans-serif;
        font-weight: normal;
        color: #333333;
        
        /* Выравнивание */
        text-align: left;
    }
    
</style>
""", unsafe_allow_html=True)
    
    # Результаты тестов
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("📊 ACT (контроль астмы)")
        act_text, act_type = interpret_act(patient['act_score'])
        if act_type == "success":
            st.success(act_text)
        elif act_type == "warning":
            st.warning(act_text)
        else:
            st.error(act_text)
        
        st.subheader("😟 HADS — Тревога")
        hads_a_text, hads_a_type = interpret_hads(patient['hads_a_score'], "anxiety")
        if hads_a_type == "success":
            st.success(hads_a_text)
        elif hads_a_type == "warning":
            st.warning(hads_a_text)
        else:
            st.error(hads_a_text)
    
    with col2:
        st.subheader("🏥 CIRS (коморбидность)")
        cirs_text, cirs_type = interpret_cirs(patient['cirs_score'])
        if cirs_type == "success":
            st.success(cirs_text)
        elif cirs_type == "warning":
            st.warning(cirs_text)
        else:
            st.error(cirs_text)
        
        st.subheader("😢 HADS — Депрессия")
        hads_d_text, hads_d_type = interpret_hads(patient['hads_d_score'], "depression")
        if hads_d_type == "success":
            st.success(hads_d_text)
        elif hads_d_type == "warning":
            st.warning(hads_d_text)
        else:
            st.error(hads_d_text)
    
    
    
    # Кнопка сохранения в Word
    st.divider()
    
    doc = create_word_report(
        patient,
        patient['act_score'], interpret_act(patient['act_score'])[0],
        patient['hads_a_score'], interpret_hads(patient['hads_a_score'], "anxiety")[0],
        patient['hads_d_score'], interpret_hads(patient['hads_d_score'], "depression")[0],
        patient['cirs_score'], interpret_cirs(patient['cirs_score'])[0],
        patient['test_date']
    )
    
    # Сохраняем в BytesIO
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.download_button(
            label="📄 Скачать отчёт (Word)",
            data=doc_bytes,
            file_name=f"Астма-тест_{patient['fio']}_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )


    with col2:
        st.download_button(
            label="📊 Скачать результаты (Excel)",
            data=excel_file,
            file_name=f"Астма-тест_{patient['fio']}_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )
    
    with col3:
        if st.button("🔄 Новый пациент", use_container_width=True):
            st.session_state.current_patient = {
                'id': None, 'fio': '', 'birth_date': None, 'gender': '',
                'test_date': None, 'act_answers': {}, 'act_score': None,
                'hads_a_answers': {}, 'hads_a_score': None,
                'hads_d_answers': {}, 'hads_d_score': None,
                'cirs_answers': {}, 'cirs_score': None, 'completed_tests': []
            }
            st.session_state.page = "patient_info"
            st.rerun()



# ==================== БОКОВАЯ ПАНЕЛЬ ====================
def render_sidebar():
    with st.sidebar:
        st.image("https://img.icons8.com/color/96/000000/lungs.png", width=80)
        st.title("Астма-тест")
        
        st.markdown("---")
        st.markdown("### 📋 Навигация")
        
        # Кнопка нового пациента
        if st.button("🆕 Новый пациент", use_container_width=True):
            st.session_state.current_patient = {
                'id': None, 'fio': '', 'birth_date': None, 'gender': '',
                'test_date': None, 'act_answers': {}, 'act_score': None,
                'hads_a_answers': {}, 'hads_a_score': None,
                'hads_d_answers': {}, 'hads_d_score': None,
                'cirs_answers': {}, 'cirs_score': None, 'completed_tests': []
            }
            st.session_state.page = "patient_info"
            st.rerun()
        
        st.markdown("---")
        
        # Кнопки тестов (активны только если заполнена информация о пациенте)
        if st.session_state.current_patient['fio']:
            col1, col2 = st.columns(2)
            with col1:
                if st.button("📋 ACT", use_container_width=True):
                    st.session_state.page = "act_test"
                    st.rerun()
            with col2:
                if st.button("📋 HADS", use_container_width=True):
                    st.session_state.page = "hads_test"
                    st.rerun()
            
            if st.button("📋 CIRS", use_container_width=True):
                st.session_state.page = "cirs_test"
                st.rerun()
            
            if len(st.session_state.current_patient['completed_tests']) == 3:
                st.success("✅ Все тесты пройдены")
                if st.button("📊 Результаты", use_container_width=True):
                    st.session_state.page = "results"
                    st.rerun()
        else:
            st.info("👆 Сначала заполните данные пациента")
        
        st.markdown("---")
        
        # Статус заполнения
        st.markdown("### 📊 Статус")
        completed = st.session_state.current_patient['completed_tests']
        st.markdown(f"✅ ACT: {'✔️' if 'ACT' in completed else '❌'}")
        st.markdown(f"✅ HADS: {'✔️' if 'HADS' in completed else '❌'}")
        st.markdown(f"✅ CIRS: {'✔️' if 'CIRS' in completed else '❌'}")
        
        st.markdown("---")
        st.caption("Астма-тест v1.0")
        st.caption("© 2026, Проект")

# ==================== ОСНОВНАЯ ФУНКЦИЯ ====================
def main():
    render_sidebar()
    
    if st.session_state.page == "patient_info":
        render_patient_info()
    elif st.session_state.page == "act_test":
        render_act_test()
    elif st.session_state.page == "hads_test":
        render_hads_test()
    elif st.session_state.page == "cirs_test":
        render_cirs_test()
    elif st.session_state.page == "results":
        render_results()

if __name__ == "__main__":
    main()

#       $env:Path += ";C:\Users\Dasha\AppData\Local\Python\pythoncore-3.14-64\Scripts"

#       pyinstaller app.spec --clean



